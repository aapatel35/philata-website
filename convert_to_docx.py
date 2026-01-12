#!/usr/bin/env python3
"""
Convert Comprehensive Report from Markdown to Word Document
"""
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level):
    """Add heading with appropriate styling"""
    heading = doc.add_heading(text, level=level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

def add_table_from_markdown(doc, lines):
    """Convert markdown table to Word table"""
    # Parse markdown table
    rows = []
    for line in lines:
        if '|' in line and '---' not in line:
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)

    if not rows:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'

    # Fill table
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # Header row styling
            if i == 0:
                set_cell_shading(cell, '1a5f7a')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()  # Add spacing after table

def convert_markdown_to_docx():
    """Main conversion function"""
    # Read markdown file
    with open('COMPREHENSIVE_WEBSITE_REPORT.md', 'r') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set document title
    title = doc.add_heading('Philata Immigration Website', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Comprehensive Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)

    doc.add_paragraph('Generated: January 11, 2026')
    doc.add_paragraph()
    doc.add_paragraph('Purpose: Complete documentation of all website sections, inputs collected, calculations performed, and outcomes provided')
    doc.add_paragraph()

    # Process content
    lines = content.split('\n')
    i = 0
    table_lines = []
    in_table = False

    while i < len(lines):
        line = lines[i]

        # Skip document header (already added)
        if line.startswith('# Philata') or line.startswith('## Document Overview') or line.startswith('**Generated:**') or line.startswith('**Purpose:**'):
            i += 1
            continue

        # Handle tables
        if '|' in line and not line.strip().startswith('|---'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and '|' not in line:
            # End of table
            add_table_from_markdown(doc, table_lines)
            in_table = False
            table_lines = []

        # Handle headings
        if line.startswith('# TABLE OF CONTENTS'):
            add_heading(doc, 'Table of Contents', 1)
            i += 1
            continue
        elif line.startswith('# ') and not line.startswith('# TABLE'):
            text = line.replace('# ', '').strip()
            add_heading(doc, text, 1)
            i += 1
            continue
        elif line.startswith('## '):
            text = line.replace('## ', '').strip()
            add_heading(doc, text, 2)
            i += 1
            continue
        elif line.startswith('### '):
            text = line.replace('### ', '').strip()
            add_heading(doc, text, 3)
            i += 1
            continue

        # Handle regular text
        if line.strip() and not line.startswith('---') and not line.startswith('['):
            # Clean markdown formatting
            text = line.strip()
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # Bold
            text = re.sub(r'\*(.*?)\*', r'\1', text)  # Italic
            text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)  # Links

            if text.startswith('- '):
                # Bullet point
                para = doc.add_paragraph(text[2:], style='List Bullet')
            elif text.startswith('1. ') or re.match(r'^\d+\.', text):
                # Numbered list
                para = doc.add_paragraph(text, style='List Number')
            else:
                para = doc.add_paragraph(text)

        i += 1

    # Handle any remaining table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)

    # Save document
    output_path = 'COMPREHENSIVE_WEBSITE_REPORT.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    convert_markdown_to_docx()
